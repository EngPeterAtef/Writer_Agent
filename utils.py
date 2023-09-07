import requests
from docx import Document
from docx.shared import Inches
import io
import os
def count_words_with_bullet_points(input_string):
    bullet_points = [
        "*",
        "-",
        "+",
        ".",
    ]  # define the bullet points to look for
    words_count = 0
    for bullet_point in bullet_points:
        input_string = input_string.replace(
            bullet_point, ""
        )  # remove the bullet points
    words_count = len(input_string.split())  # count the words
    return words_count

def get_src_original_url(query):
    url = "https://api.pexels.com/v1/search"
    headers = {
        "Authorization": os.getenv("PIXELES_API_KEY"),
    }

    params = {
        "query": query,
        "per_page": 1,
    }

    response = requests.get(url, headers=headers, params=params)

    # Check if the request was successful (status code 200)
    if response.status_code == 200:
        data = response.json()
        print("data",data)
        photos = data.get("photos", [])
        print("photos",photos)
        if photos != []:
            src_original_url = photos[0]["src"]["original"]
            return src_original_url
        else:
            return ""
    else:
        return None


def create_word_docx(user_input, paragraph, image_input):
    # Create a new Word document
    doc = Document()

    # Add the user input to the document
    doc.add_heading(user_input, level=1)
    doc.add_paragraph(paragraph)

    # Add the image to the document
    doc.add_heading("Image Input", level=1)
    image_stream = io.BytesIO()
    image_input.save(image_stream, format="PNG")
    image_stream.seek(0)
    doc.add_picture(image_stream, width=Inches(4))  # Adjust the width as needed
    return doc